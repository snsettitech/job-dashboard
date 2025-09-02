# resume-service/app/services/file_processor.py
import fitz  # PyMuPDF for PDF processing
from docx import Document  # python-docx for DOCX processing
import io
from typing import Optional, Dict, Any, List
import re
import logging
from PIL import Image
import pytesseract
import cv2
import numpy as np
from datetime import datetime

logger = logging.getLogger(__name__)

class FileProcessor:
    """Enhanced file processor with cloud storage integration"""
    
    def __init__(self):
        self.supported_types = {
            "application/pdf": self._extract_from_pdf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_from_docx,
            "text/plain": self._extract_from_txt,
            "image/jpeg": self._extract_from_image,
            "image/png": self._extract_from_image,
            "image/jpg": self._extract_from_image
        }
    
    async def extract_text(self, content: bytes, content_type: str, filename: str = "") -> Dict[str, Any]:
        """Extract text from uploaded file with enhanced processing"""
        try:
            if content_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Extract text using appropriate method
            extracted_text = self.supported_types[content_type](content)
            
            # Clean and validate text
            cleaned_text = self._clean_text(extracted_text)
            
            if not cleaned_text.strip():
                raise ValueError("No text could be extracted from the file")
            
            # Extract structured information
            structured_info = self._extract_structured_info(cleaned_text)
            
            # Quality assessment
            quality_assessment = self._assess_quality(cleaned_text, structured_info)
            
            return {
                "raw_text": extracted_text,
                "cleaned_text": cleaned_text,
                "structured_info": structured_info,
                "quality_assessment": quality_assessment,
                "metadata": {
                    "filename": filename,
                    "content_type": content_type,
                    "file_size": len(content),
                    "word_count": len(cleaned_text.split()),
                    "character_count": len(cleaned_text),
                    "processing_timestamp": datetime.utcnow().isoformat()
                }
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from file: {e}")
            raise ValueError(f"Failed to process file: {str(e)}")
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF with OCR fallback"""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # If no text found, try OCR
                if not page_text.strip():
                    page_text = self._ocr_page(page)
                
                text += page_text + "\n"
            
            doc.close()
            return text
            
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX with enhanced table handling"""
        try:
            doc = Document(io.BytesIO(content))
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables with better formatting
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"  # Add spacing between tables
            
            return text
            
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {str(e)}")
    
    def _extract_from_txt(self, content: bytes) -> str:
        """Extract text from plain text file with encoding detection"""
        try:
            encodings = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1', 'latin-1']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise ValueError(f"Error processing text file: {str(e)}")
    
    def _extract_from_image(self, content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            # Convert bytes to PIL Image
            image = Image.open(io.BytesIO(content))
            
            # Convert to OpenCV format
            cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Preprocess image for better OCR
            processed_image = self._preprocess_image_for_ocr(cv_image)
            
            # Perform OCR
            text = pytesseract.image_to_string(processed_image)
            
            return text
            
        except Exception as e:
            raise ValueError(f"Error processing image: {str(e)}")
    
    def _preprocess_image_for_ocr(self, image):
        """Preprocess image for better OCR results"""
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply thresholding
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Remove noise
        kernel = np.ones((1, 1), np.uint8)
        opening = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel)
        
        return opening
    
    def _ocr_page(self, page) -> str:
        """Perform OCR on a PDF page"""
        try:
            # Convert page to image
            pix = page.get_pixmap()
            img_data = pix.tobytes("png")
            
            # Convert to PIL Image
            image = Image.open(io.BytesIO(img_data))
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            return text
            
        except Exception as e:
            logger.warning(f"OCR failed for page: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Enhanced text cleaning and normalization"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep necessary punctuation
        text = re.sub(r'[^\w\s\.,;:()\-@/]', '', text)
        
        # Fix common OCR/extraction errors
        text = text.replace('•', '-')  # Replace bullets
        text = text.replace('  ', ' ')  # Double spaces
        text = text.replace('\n\n\n', '\n\n')  # Multiple newlines
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def _extract_structured_info(self, text: str) -> Dict[str, Any]:
        """Extract structured information from resume text"""
        info = {
            "contact_info": self._extract_contact_info(text),
            "skills": self._extract_skills_section(text),
            "experience": self._extract_experience_section(text),
            "education": self._extract_education_section(text),
            "sections": self._identify_sections(text),
            "achievements": self._extract_achievements(text)
        }
        
        return info
    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information with enhanced patterns"""
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "website": None
        }
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group()
        
        # Phone patterns (multiple formats)
        phone_patterns = [
            r'(\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
            r'\+?[\d\s\-\(\)]{10,}',  # International format
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}'  # Simple format
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact_info["phone"] = phone_match.group()
                break
        
        # LinkedIn profile
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text.lower())
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group()
        
        # GitHub profile
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text.lower())
        if github_match:
            contact_info["github"] = github_match.group()
        
        # Website
        website_pattern = r'https?://(?:www\.)?[a-zA-Z0-9-]+\.[a-zA-Z]{2,}(?:/\S*)?'
        website_match = re.search(website_pattern, text)
        if website_match:
            contact_info["website"] = website_match.group()
        
        return contact_info
    
    def _extract_skills_section(self, text: str) -> Dict[str, Any]:
        """Extract skills section with enhanced skill detection"""
        skills_keywords = ['skills', 'technical skills', 'technologies', 'programming languages', 'tools', 'competencies']
        
        # Find skills section
        skills_text = ""
        lines = text.lower().split('\n')
        
        in_skills_section = False
        for i, line in enumerate(lines):
            if any(keyword in line for keyword in skills_keywords):
                in_skills_section = True
                continue
            elif in_skills_section and (line.strip() == '' or any(keyword in line for keyword in ['experience', 'education', 'projects', 'work history'])):
                break
            elif in_skills_section:
                skills_text += line + " "
        
        # Enhanced skill detection
        common_skills = [
            # Programming languages
            'python', 'javascript', 'java', 'c++', 'c#', 'go', 'rust', 'typescript', 'php', 'ruby', 'swift', 'kotlin', 'scala',
            # Web technologies
            'react', 'angular', 'vue', 'node.js', 'express', 'fastapi', 'django', 'flask', 'spring', 'html', 'css', 'sass', 'less',
            # Databases
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'sqlite', 'oracle', 'sql server',
            # Cloud/DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'git', 'github', 'gitlab', 'ci/cd',
            # AI/ML
            'machine learning', 'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy', 'matplotlib', 'seaborn',
            # Other
            'linux', 'api', 'microservices', 'agile', 'scrum', 'jira', 'confluence', 'slack', 'figma', 'adobe'
        ]
        
        found_skills = []
        text_lower = (skills_text + " " + text).lower()
        
        for skill in common_skills:
            if skill in text_lower:
                found_skills.append(skill)
        
        return {
            "raw_skills_text": skills_text.strip(),
            "identified_skills": found_skills,
            "skill_count": len(found_skills),
            "skill_categories": self._categorize_skills(found_skills)
        }
    
    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize skills by type"""
        categories = {
            "programming_languages": [],
            "frameworks": [],
            "databases": [],
            "cloud_platforms": [],
            "tools": [],
            "methodologies": []
        }
        
        skill_categories = {
            "programming_languages": ['python', 'javascript', 'java', 'c++', 'c#', 'go', 'rust', 'typescript', 'php', 'ruby', 'swift', 'kotlin', 'scala'],
            "frameworks": ['react', 'angular', 'vue', 'express', 'fastapi', 'django', 'flask', 'spring'],
            "databases": ['postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'sqlite', 'oracle', 'sql server'],
            "cloud_platforms": ['aws', 'azure', 'gcp'],
            "tools": ['docker', 'kubernetes', 'terraform', 'jenkins', 'git', 'github', 'gitlab', 'jira', 'figma'],
            "methodologies": ['agile', 'scrum', 'microservices', 'ci/cd']
        }
        
        for skill in skills:
            for category, category_skills in skill_categories.items():
                if skill in category_skills:
                    categories[category].append(skill)
                    break
        
        return categories
    
    def _extract_experience_section(self, text: str) -> Dict[str, Any]:
        """Extract work experience information"""
        experience_keywords = ['experience', 'work history', 'employment', 'professional experience', 'work experience']
        
        # Enhanced experience detection
        experience_text = ""
        lines = text.lower().split('\n')
        
        in_experience_section = False
        for i, line in enumerate(lines):
            if any(keyword in line for keyword in experience_keywords):
                in_experience_section = True
                continue
            elif in_experience_section and (line.strip() == '' or any(keyword in line for keyword in ['education', 'projects', 'skills'])):
                break
            elif in_experience_section:
                experience_text += line + " "
        
        return {
            "has_experience_section": any(keyword in text.lower() for keyword in experience_keywords),
            "years_pattern_found": bool(re.search(r'\d+\+?\s*years?', text.lower())),
            "company_patterns": len(re.findall(r'\b(?:inc|llc|corp|company|ltd|tech|solutions|systems)\b', text.lower())),
            "experience_text": experience_text.strip(),
            "job_titles": self._extract_job_titles(text)
        }
    
    def _extract_job_titles(self, text: str) -> List[str]:
        """Extract job titles from text"""
        job_title_patterns = [
            r'(?:senior|junior|lead|principal|staff)?\s*(?:software engineer|developer|programmer|architect|manager|director|consultant)',
            r'(?:full stack|front end|back end|devops|data|machine learning|ai|ml)\s*(?:engineer|developer)',
            r'(?:product|project|program|engineering)\s*(?:manager|director|lead)'
        ]
        
        job_titles = []
        for pattern in job_title_patterns:
            matches = re.findall(pattern, text.lower())
            job_titles.extend(matches)
        
        return list(set(job_titles))
    
    def _extract_education_section(self, text: str) -> Dict[str, Any]:
        """Extract education information"""
        education_keywords = ['education', 'degree', 'university', 'college', 'bachelor', 'master', 'phd']
        degrees = ['bachelor', 'master', 'phd', 'doctorate', 'associates', 'b.s.', 'm.s.', 'b.a.', 'm.a.', 'b.tech', 'm.tech']
        
        return {
            "has_education_section": any(keyword in text.lower() for keyword in education_keywords),
            "degrees_mentioned": [degree for degree in degrees if degree in text.lower()],
            "graduation_years": re.findall(r'\b(19|20)\d{2}\b', text),
            "universities": self._extract_universities(text)
        }
    
    def _extract_universities(self, text: str) -> List[str]:
        """Extract university names from text"""
        university_keywords = ['university', 'college', 'institute', 'school']
        universities = []
        
        lines = text.split('\n')
        for line in lines:
            if any(keyword in line.lower() for keyword in university_keywords):
                universities.append(line.strip())
        
        return universities
    
    def _extract_achievements(self, text: str) -> List[str]:
        """Extract achievements and accomplishments"""
        achievement_patterns = [
            r'(?:increased|improved|reduced|achieved|delivered|led|managed|developed|created|implemented).*?(?:by \d+%|\d+%|\d+x|\d+ times)',
            r'(?:awarded|recognized|honored|received).*?(?:award|recognition|honor)',
            r'(?:successfully|effectively|efficiently).*?(?:completed|delivered|implemented)'
        ]
        
        achievements = []
        for pattern in achievement_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            achievements.extend(matches)
        
        return achievements
    
    def _identify_sections(self, text: str) -> List[str]:
        """Identify main resume sections"""
        common_sections = [
            'summary', 'objective', 'skills', 'experience', 'education', 'projects', 
            'certifications', 'awards', 'publications', 'languages', 'volunteer',
            'interests', 'references'
        ]
        
        found_sections = []
        text_lower = text.lower()
        
        for section in common_sections:
            if section in text_lower:
                found_sections.append(section)
        
        return found_sections
    
    def _assess_quality(self, text: str, structured_info: Dict[str, Any]) -> Dict[str, Any]:
        """Assess resume quality and completeness"""
        score = 0
        max_score = 100
        feedback = []
        
        # Check contact information (20 points)
        contact_info = structured_info["contact_info"]
        if contact_info["email"]:
            score += 10
        else:
            feedback.append("Add an email address")
        
        if contact_info["phone"]:
            score += 5
        else:
            feedback.append("Add a phone number")
        
        if contact_info["linkedin"]:
            score += 5
        else:
            feedback.append("Add LinkedIn profile")
        
        # Check skills section (25 points)
        skills_info = structured_info["skills"]
        if skills_info["skill_count"] > 0:
            score += min(25, skills_info["skill_count"] * 3)
        else:
            feedback.append("Add a skills section with relevant technical skills")
        
        # Check experience section (30 points)
        experience_info = structured_info["experience"]
        if experience_info["has_experience_section"]:
            score += 20
            if experience_info["years_pattern_found"]:
                score += 5
            if experience_info["company_patterns"] > 0:
                score += 5
        else:
            feedback.append("Add work experience section")
        
        # Check education section (15 points)
        education_info = structured_info["education"]
        if education_info["has_education_section"]:
            score += 15
        else:
            feedback.append("Add education section")
        
        # Check overall structure (10 points)
        sections = structured_info["sections"]
        if len(sections) >= 4:
            score += 10
        else:
            feedback.append(f"Add more sections (found {len(sections)}, recommended 4+)")
        
        # Determine grade
        if score >= 90:
            grade = "Excellent"
        elif score >= 75:
            grade = "Good"
        elif score >= 60:
            grade = "Fair"
        else:
            grade = "Needs Improvement"
        
        return {
            "quality_score": min(score, max_score),
            "max_score": max_score,
            "feedback": feedback,
            "grade": grade,
            "completeness_score": score / max_score * 100
        }

