# backend/app/services/document_generator.py
import io
import tempfile
import os
from typing import Dict, Any, Optional
from datetime import datetime
import logging

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import black, blue, darkblue, white
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
    print("ReportLab imported successfully")
except ImportError as e:
    REPORTLAB_AVAILABLE = False
    logging.warning(f"ReportLab not available - PDF generation disabled: {e}")
    # Define dummy classes to prevent import errors
    class Table:
        def __init__(self, *args, **kwargs):
            pass
        def setStyle(self, *args, **kwargs):
            pass
    class TableStyle:
        def __init__(self, *args, **kwargs):
            pass

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX generation disabled")

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Service for generating professional PDF and DOCX documents from resume text"""
    
    def __init__(self):
        self.page_size = A4
        self.margin = 0.75 * inch
        
    def generate_pdf(self, resume_text: str, filename: str = "resume.pdf") -> bytes:
        """Generate a professional PDF resume"""
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("ReportLab not available for PDF generation")
        
        try:
            # Create PDF document
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=self.page_size, 
                                 leftMargin=self.margin, rightMargin=self.margin,
                                 topMargin=self.margin, bottomMargin=self.margin)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12,
                alignment=TA_CENTER,
                textColor=darkblue
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=8,
                spaceBefore=12,
                textColor=blue
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                alignment=TA_LEFT
            )
            
            # Parse resume text into sections
            sections = self._parse_resume_sections(resume_text)
            
            # Build story
            story = []
            
            # Add header
            if 'name' in sections:
                story.append(Paragraph(sections['name'], title_style))
                story.append(Spacer(1, 12))
            
            # Add contact info
            if 'contact' in sections:
                contact_table = self._create_contact_table(sections['contact'])
                story.append(contact_table)
                story.append(Spacer(1, 12))
            
            # Add sections
            for section_name, content in sections.items():
                if section_name in ['name', 'contact']:
                    continue
                    
                if content.strip():
                    story.append(Paragraph(section_name.replace('_', ' ').title(), heading_style))
                    
                    # Split content into paragraphs
                    paragraphs = content.strip().split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            story.append(Paragraph(para.strip(), normal_style))
                    
                    story.append(Spacer(1, 8))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise RuntimeError(f"Failed to generate PDF: {str(e)}")
    
    def generate_docx(self, resume_text: str, filename: str = "resume.docx") -> bytes:
        """Generate a professional DOCX resume"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available for DOCX generation")
        
        try:
            # Create document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Parse resume text into sections
            sections = self._parse_resume_sections(resume_text)
            
            # Add header
            if 'name' in sections:
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(sections['name'])
                name_run.font.size = Pt(18)
                name_run.font.bold = True
                name_run.font.color.rgb = doc.styles['Normal'].font.color
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Add spacing
            
            # Add contact info
            if 'contact' in sections:
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_run = contact_para.add_run(sections['contact'])
                contact_run.font.size = Pt(11)
                doc.add_paragraph()  # Add spacing
            
            # Add sections
            for section_name, content in sections.items():
                if section_name in ['name', 'contact']:
                    continue
                    
                if content.strip():
                    # Add section heading
                    heading_para = doc.add_paragraph()
                    heading_run = heading_para.add_run(section_name.replace('_', ' ').title())
                    heading_run.font.size = Pt(14)
                    heading_run.font.bold = True
                    heading_run.font.color.rgb = doc.styles['Normal'].font.color
                    
                    # Add section content
                    paragraphs = content.strip().split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                    
                    doc.add_paragraph()  # Add spacing
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise RuntimeError(f"Failed to generate DOCX: {str(e)}")
    
    def _parse_resume_sections(self, resume_text: str) -> Dict[str, str]:
        """Parse resume text into logical sections"""
        sections = {}
        
        # Split text into lines
        lines = resume_text.split('\n')
        
        current_section = 'name'
        current_content = []
        
        for line in lines:
            line = line.strip()
            
            # Detect section headers (common resume section names)
            section_keywords = [
                'education', 'experience', 'work', 'employment', 'skills', 'technical',
                'projects', 'achievements', 'certifications', 'languages', 'interests',
                'summary', 'objective', 'profile', 'qualifications', 'publications',
                'presentations', 'awards', 'honors', 'volunteer', 'activities'
            ]
            
            # Check if line is a section header
            is_header = False
            for keyword in section_keywords:
                if line.lower().startswith(keyword) and len(line) < 50:
                    is_header = True
                    break
            
            if is_header and current_content:
                # Save previous section
                sections[current_section] = '\n'.join(current_content).strip()
                current_section = line.lower().replace(' ', '_')
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        # Extract name and contact info
        if 'name' in sections:
            name_lines = sections['name'].split('\n')
            if len(name_lines) > 1:
                sections['name'] = name_lines[0].strip()
                sections['contact'] = '\n'.join(name_lines[1:]).strip()
        
        return sections
    
    def _create_contact_table(self, contact_info: str) -> Table:
        """Create a formatted contact information table"""
        contact_lines = contact_info.split('\n')
        contact_data = []
        
        for line in contact_lines:
            if line.strip():
                contact_data.append([line.strip()])
        
        if not contact_data:
            return Table([['Contact information not available']])
        
        table = Table(contact_data, colWidths=[4*inch])
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (-1, -1), black),
            ('GRID', (0, 0), (-1, -1), 0, white),  # No grid
            ('BACKGROUND', (0, 0), (-1, -1), white),
        ]))
        
        return table
