# backend/app/services/file_processor.py - Resume File Processing
import fitz  # PyMuPDF for PDF processing
from docx import Document  # python-docx for DOCX processing
import io
from typing import Optional, Dict, Any, List
import re
import logging

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles extraction of text from various file formats"""
    
    def __init__(self):
        self.supported_types = {
            "application/pdf": self._extract_from_pdf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_from_docx,
            "text/plain": self._extract_from_txt
        }
    
    async def extract_text(self, content: bytes, content_type: str, filename: str = "") -> Dict[str, Any]:
        """
        Extract text from uploaded file
        Returns dict with extracted text and metadata
        """
        try:
            if content_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Extract text using appropriate method
            extracted_text = self.supported_types[content_type](content)
            
            # Clean and validate text
            cleaned_text = self._clean_text(extracted_text)
            
            if not cleaned_text.strip():
                raise ValueError("No text could be extracted from the file")
            
            # Extract structured information
            structured_info = self._extract_structured_info(cleaned_text)
            
            return {
                "raw_text": extracted_text,
                "cleaned_text": cleaned_text,
                "structured_info": structured_info,
                "metadata": {
                    "filename": filename,
                    "content_type": content_type,
                    "file_size": len(content),
                    "word_count": len(cleaned_text.split()),
                    "character_count": len(cleaned_text)
                }
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from file: {e}")
            raise ValueError(f"Failed to process file: {str(e)}")
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
                text += "\n"  # Add page break
            
            doc.close()
            return text
            
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            doc = Document(io.BytesIO(content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text
            
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {str(e)}")
    
    def _extract_from_txt(self, content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise ValueError(f"Error processing text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep necessary punctuation
        text = re.sub(r'[^\w\s\.,;:()\-@/]', '', text)
        
        # Fix common OCR/extraction errors
        text = text.replace('•', '-')  # Replace bullets
        text = text.replace('  ', ' ')  # Double spaces
        
        return text.strip()
    
    def _extract_structured_info(self, text: str) -> Dict[str, Any]:
        """Extract structured information from resume text"""
        info = {
            "contact_info": self._extract_contact_info(text),
            "skills": self._extract_skills_section(text),
            "experience": self._extract_experience_section(text),
            "education": self._extract_education_section(text),
            "sections": self._identify_sections(text)
        }
        
        return info
    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information"""
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None
        }
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group()
        
        # Phone pattern (US format)
        phone_pattern = r'(\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info["phone"] = phone_match.group()
        
        # LinkedIn profile
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text.lower())
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group()
        
        # GitHub profile
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text.lower())
        if github_match:
            contact_info["github"] = github_match.group()
        
        return contact_info
    
    def _extract_skills_section(self, text: str) -> Dict[str, Any]:
        """Extract skills section and categorize skills"""
        skills_keywords = ['skills', 'technical skills', 'technologies', 'programming languages', 'tools']
        
        # Find skills section
        skills_text = ""
        lines = text.lower().split('\n')
        
        in_skills_section = False
        for i, line in enumerate(lines):
            if any(keyword in line for keyword in skills_keywords):
                in_skills_section = True
                continue
            elif in_skills_section and (line.strip() == '' or any(keyword in line for keyword in ['experience', 'education', 'projects'])):
                break
            elif in_skills_section:
                skills_text += line + " "
        
        # Extract individual skills
        common_skills = [
            # Programming languages
            'python', 'javascript', 'java', 'c++', 'c#', 'go', 'rust', 'typescript', 'php', 'ruby', 'swift', 'kotlin',
            # Web technologies
            'react', 'angular', 'vue', 'node.js', 'express', 'fastapi', 'django', 'flask', 'spring', 'html', 'css',
            # Databases
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'sqlite',
            # Cloud/DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'git',
            # AI/ML
            'machine learning', 'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy',
            # Other
            'linux', 'api', 'microservices', 'agile', 'scrum'
        ]
        
        found_skills = []
        text_lower = (skills_text + " " + text).lower()
        
        for skill in common_skills:
            if skill in text_lower:
                found_skills.append(skill)
        
        return {
            "raw_skills_text": skills_text.strip(),
            "identified_skills": found_skills,
            "skill_count": len(found_skills)
        }
    
    def _extract_experience_section(self, text: str) -> Dict[str, Any]:
        """Extract work experience information"""
        experience_keywords = ['experience', 'work history', 'employment', 'professional experience']
        
        # This is a simplified version - could be enhanced with NLP
        return {
            "has_experience_section": any(keyword in text.lower() for keyword in experience_keywords),
            "years_pattern_found": bool(re.search(r'\d+\+?\s*years?', text.lower())),
            "company_patterns": len(re.findall(r'\b(?:inc|llc|corp|company|ltd)\b', text.lower()))
        }
    
    def _extract_education_section(self, text: str) -> Dict[str, Any]:
        """Extract education information"""
        education_keywords = ['education', 'degree', 'university', 'college', 'bachelor', 'master', 'phd']
        degrees = ['bachelor', 'master', 'phd', 'doctorate', 'associates', 'b.s.', 'm.s.', 'b.a.', 'm.a.']
        
        return {
            "has_education_section": any(keyword in text.lower() for keyword in education_keywords),
            "degrees_mentioned": [degree for degree in degrees if degree in text.lower()],
            "graduation_years": re.findall(r'\b(19|20)\d{2}\b', text)
        }
    
    def _identify_sections(self, text: str) -> List[str]:
        """Identify main resume sections"""
        common_sections = [
            'summary', 'objective', 'skills', 'experience', 'education', 'projects', 
            'certifications', 'awards', 'publications', 'languages'
        ]
        
        found_sections = []
        text_lower = text.lower()
        
        for section in common_sections:
            if section in text_lower:
                found_sections.append(section)
        
        return found_sections

# Utility function to validate resume quality
def validate_resume_quality(extracted_info: Dict[str, Any]) -> Dict[str, Any]:
    """Validate resume completeness and quality"""
    
    score = 0
    max_score = 100
    feedback = []
    
    # Check contact information (20 points)
    contact_info = extracted_info["structured_info"]["contact_info"]
    if contact_info["email"]:
        score += 10
    else:
        feedback.append("Add an email address")
    
    if contact_info["phone"]:
        score += 5
    else:
        feedback.append("Add a phone number")
    
    if contact_info["linkedin"]:
        score += 5
    else:
        feedback.append("Add LinkedIn profile")
    
    # Check skills section (25 points)
    skills_info = extracted_info["structured_info"]["skills"]
    if skills_info["skill_count"] > 0:
        score += min(25, skills_info["skill_count"] * 3)
    else:
        feedback.append("Add a skills section with relevant technical skills")
    
    # Check experience section (30 points)
    experience_info = extracted_info["structured_info"]["experience"]
    if experience_info["has_experience_section"]:
        score += 20
        if experience_info["years_pattern_found"]:
            score += 5
        if experience_info["company_patterns"] > 0:
            score += 5
    else:
        feedback.append("Add work experience section")
    
    # Check education section (15 points)
    education_info = extracted_info["structured_info"]["education"]
    if education_info["has_education_section"]:
        score += 15
    else:
        feedback.append("Add education section")
    
    # Check overall structure (10 points)
    sections = extracted_info["structured_info"]["sections"]
    if len(sections) >= 4:
        score += 10
    else:
        feedback.append(f"Add more sections (found {len(sections)}, recommended 4+)")
    
    return {
        "quality_score": min(score, max_score),
        "max_score": max_score,
        "feedback": feedback,
        "grade": "Excellent" if score >= 90 else "Good" if score >= 75 else "Fair" if score >= 60 else "Needs Improvement"
    }